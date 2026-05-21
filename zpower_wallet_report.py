#!/usr/bin/env python3
"""
ZPower Ethereum Wallet Forensics Tool
======================================
Usage:
    python zpower_wallet_report.py

Requirements:
    pip install requests pandas python-docx openpyxl
"""

import requests
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timezone
import time, sys

BASE_V2 = "https://api.etherscan.io/v2/api"
CHAINID = "1"

def api_call(action, module, extra_params, api_key):
    params = {"chainid": CHAINID, "module": module, "action": action, "apikey": api_key}
    params.update(extra_params)
    while True:
        r = requests.get(BASE_V2, params=params, timeout=30)
        data = r.json()
        if data.get("message") == "OK" or data.get("status") == "1":
            return data.get("result", [])
        if "rate limit" in str(data).lower():
            time.sleep(1.2)
            continue
        return [] if data.get("result") == "0" else data.get("result", [])

def get_eth_price(api_key):
    params = {"chainid": CHAINID, "module": "stats", "action": "ethprice", "apikey": api_key}
    r = requests.get(BASE_V2, params=params, timeout=15).json()
    try:
        return float(r["result"]["ethusd"])
    except Exception:
        return 0.0

def get_balance(address, api_key):
    result = api_call("balance", "account", {"address": address, "tag": "latest"}, api_key)
    if isinstance(result, str):
        try:
            return int(result) / 1e18
        except Exception:
            return 0.0
    return 0.0

def get_all_txs(address, api_key):
    txs = []
    for page in range(1, 50):
        batch = api_call("txlist", "account", {
            "address": address, "startblock": 0, "endblock": 99999999,
            "sort": "asc", "page": page, "offset": 1000,
        }, api_key)
        if not batch:
            break
        txs.extend(batch)
        if len(batch) < 1000:
            break
        time.sleep(0.3)
    return txs

def get_all_token_txs(address, api_key):
    txs = []
    for page in range(1, 20):
        batch = api_call("tokentx", "account", {
            "address": address, "startblock": 0, "endblock": 99999999,
            "sort": "asc", "page": page, "offset": 1000,
        }, api_key)
        if not batch:
            break
        txs.extend(batch)
        if len(batch) < 1000:
            break
        time.sleep(0.3)
    return txs

def process_txs(address, txs, eth_price):
    address_lower = address.lower()
    rows = []
    for tx in txs:
        if tx.get("isError", "0") == "1":
            continue
        try:
            value_eth = int(tx.get("value", 0)) / 1e18
        except Exception:
            value_eth = 0.0
        ts = int(tx.get("timeStamp", 0))
        dt = datetime.fromtimestamp(ts, tz=timezone.utc) if ts else None
        fr = tx.get("from", "").lower()
        to = tx.get("to", "").lower()
        direction = "IN" if to == address_lower else "OUT"
        counterparty = fr if direction == "IN" else to
        rows.append({
            "hash": tx.get("hash", ""), "timestamp": dt, "direction": direction,
            "counterparty": counterparty, "value_eth": value_eth,
            "value_usd": value_eth * eth_price, "token": "ETH",
            "gas_used": int(tx.get("gasUsed", 0)),
        })
    return rows

def process_token_txs(address, token_txs):
    address_lower = address.lower()
    rows = []
    for tx in token_txs:
        try:
            decimals = int(tx.get("tokenDecimal", 18))
            value = int(tx.get("value", 0)) / (10 ** decimals)
        except Exception:
            value = 0.0
        ts = int(tx.get("timeStamp", 0))
        dt = datetime.fromtimestamp(ts, tz=timezone.utc) if ts else None
        fr = tx.get("from", "").lower()
        to = tx.get("to", "").lower()
        direction = "IN" if to == address_lower else "OUT"
        counterparty = fr if direction == "IN" else to
        rows.append({
            "hash": tx.get("hash", ""), "timestamp": dt, "direction": direction,
            "counterparty": counterparty, "value_eth": 0.0, "value_usd": value,
            "token": tx.get("tokenSymbol", "?"), "gas_used": 0,
        })
    return rows

def top_counterparties(df, direction, n=10):
    sub = df[df["direction"] == direction].copy()
    agg = (
        sub.groupby("counterparty")
        .agg(tx_count=("hash", "count"), total_eth=("value_eth", "sum"),
             total_usd=("value_usd", "sum"),
             tokens=("token", lambda x: ", ".join(sorted(set(x)))))
        .reset_index()
        .sort_values("total_usd", ascending=False)
        .head(n)
    )
    return agg

def set_cell_color(cell, color_rgb):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    r, g, b = color_rgb
    shd.set(qn("w:fill"), f"{r:02X}{g:02X}{b:02X}")
    tcPr.append(shd)

def build_docx(address, balance_eth, eth_price, df, output_path):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("ZPower Ethereum Wallet Forensic Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 70, 127)
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    doc.add_heading("1. Wallet Summary", level=2)
    balance_usd = balance_eth * eth_price
    eth_txs = df[df["token"] == "ETH"]
    all_in = df[df["direction"] == "IN"]
    all_out = df[df["direction"] == "OUT"]
    in_eth = eth_txs[eth_txs["direction"] == "IN"]["value_eth"].sum()
    out_eth = eth_txs[eth_txs["direction"] == "OUT"]["value_eth"].sum()
    first_ts = df["timestamp"].min()
    wallet_age_days = (datetime.now(tz=timezone.utc) - first_ts).days if first_ts else 0

    summary_data = [
        ("Wallet Address", address),
        ("Etherscan Link", f"https://etherscan.io/address/{address}"),
        ("Wallet Age", f"{wallet_age_days // 365}y {wallet_age_days % 365}d (since {first_ts.strftime('%Y-%m-%d') if first_ts else 'N/A'})"),
        ("ETH Balance", f"{balance_eth:.6f} ETH  =  ${balance_usd:,.2f} USD"),
        ("ETH Price", f"${eth_price:,.2f} USD"),
        ("Total Transactions", f"{len(df):,}"),
        ("Received (IN)", f"{len(all_in):,} txs | +{in_eth:.4f} ETH | +${all_in['value_usd'].sum():,.2f} USD"),
        ("Sent (OUT)", f"{len(all_out):,} txs | -{out_eth:.4f} ETH | -${all_out['value_usd'].sum():,.2f} USD"),
        ("Net ETH Flow", f"{in_eth - out_eth:.4f} ETH"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_color(cell, (68, 114, 196))
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    for metric, value in summary_data:
        row = table.add_row()
        row.cells[0].text = metric
        row.cells[1].text = value

    doc.add_paragraph("")
    doc.add_heading("2. Token Breakdown", level=2)
    token_agg = (df.groupby(["token", "direction"])
        .agg(tx_count=("hash", "count"), total_val=("value_usd", "sum"))
        .reset_index().sort_values("total_val", ascending=False))
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    for i, h in enumerate(["Token", "Direction", "Tx Count", "Total Value (USD)"]):
        t2.rows[0].cells[i].text = h
        set_cell_color(t2.rows[0].cells[i], (68, 114, 196))
        t2.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for _, rd in token_agg.iterrows():
        row = t2.add_row()
        sign = "+" if rd["direction"] == "IN" else "-"
        for i, v in enumerate([rd["token"], rd["direction"], f"{rd['tx_count']:,}", f"{sign}${rd['total_val']:,.2f}"]):
            row.cells[i].text = str(v)
        color = (198, 239, 206) if rd["direction"] == "IN" else (255, 199, 206)
        for cell in row.cells:
            set_cell_color(cell, color)

    doc.add_paragraph("")
    for direction, label, sign in [("IN", "3a. Top 10 Incoming Wallets", "+"), ("OUT", "3b. Top 10 Outgoing Wallets", "-")]:
        doc.add_heading(label, level=2)
        top = top_counterparties(df, direction, 10)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        for i, h in enumerate(["#", "Counterparty Address", "Tx Count", "Total ETH", "Total USD"]):
            t.rows[0].cells[i].text = h
            set_cell_color(t.rows[0].cells[i], (68, 114, 196))
            t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        color = (198, 239, 206) if direction == "IN" else (255, 199, 206)
        for rank, (_, cp) in enumerate(top.iterrows(), 1):
            row = t.add_row()
            for i, v in enumerate([rank, cp["counterparty"], cp["tx_count"],
                                    f"{sign}{cp['total_eth']:.4f} ETH", f"{sign}${cp['total_usd']:,.2f}"]):
                row.cells[i].text = str(v)
                set_cell_color(row.cells[i], color)
        doc.add_paragraph("")

    doc.save(output_path)
    print(f"DOCX saved: {output_path}")

def build_excel(address, df, output_path):
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        df_out = df.copy()
        df_out["timestamp"] = df_out["timestamp"].astype(str)
        df_out["etherscan_link"] = df_out["hash"].apply(lambda h: f"https://etherscan.io/tx/{h}")
        df_out.to_excel(writer, sheet_name="All Transactions", index=False)
        top_in = top_counterparties(df, "IN", 10)
        top_in["link"] = top_in["counterparty"].apply(lambda a: f"https://etherscan.io/address/{a}")
        top_in.to_excel(writer, sheet_name="Top 10 IN", index=False)
        top_out = top_counterparties(df, "OUT", 10)
        top_out["link"] = top_out["counterparty"].apply(lambda a: f"https://etherscan.io/address/{a}")
        top_out.to_excel(writer, sheet_name="Top 10 OUT", index=False)
        token_summary = (df.groupby(["token", "direction"])
            .agg(tx_count=("hash", "count"), total_usd=("value_usd", "sum"), total_eth=("value_eth", "sum"))
            .reset_index())
        token_summary.to_excel(writer, sheet_name="Token Summary", index=False)
    print(f"Excel saved: {output_path}")

def main():
    print("=" * 60)
    print("  ZPower Ethereum Wallet Forensics Tool")
    print("=" * 60)
    address = input("\nEnter Ethereum wallet address: ").strip()
    if not address.startswith("0x") or len(address) != 42:
        print("Invalid address.")
        sys.exit(1)
    api_key = input("Enter Etherscan API key (free at https://etherscan.io/apis): ").strip()
    if not api_key:
        print("API key required.")
        sys.exit(1)

    print("\nFetching data from Etherscan...")
    eth_price = get_eth_price(api_key)
    print(f"   ETH Price: ${eth_price:,.2f}")
    balance_eth = get_balance(address, api_key)
    print(f"   Balance: {balance_eth:.6f} ETH")
    print("   Fetching normal transactions...")
    normal_txs = get_all_txs(address, api_key)
    print(f"   Found {len(normal_txs)} normal txs")
    print("   Fetching ERC-20 token transfers...")
    token_txs = get_all_token_txs(address, api_key)
    print(f"   Found {len(token_txs)} token txs")

    rows = process_txs(address, normal_txs, eth_price)
    rows += process_token_txs(address, token_txs)
    df = pd.DataFrame(rows)

    if df.empty:
        print("No transactions found.")
        sys.exit(0)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = f"wallet_report_{address[:8]}_{ts}.docx"
    excel_path = f"wallet_report_{address[:8]}_{ts}.xlsx"

    print("\nGenerating DOCX...")
    build_docx(address, balance_eth, eth_price, df, docx_path)
    print("Generating Excel...")
    build_excel(address, df, excel_path)

    print("\n" + "=" * 60)
    print(f"  DOCX:  {docx_path}")
    print(f"  Excel: {excel_path}")
    print(f"  Link:  https://etherscan.io/address/{address}")
    print("=" * 60)

if __name__ == "__main__":
    main()
